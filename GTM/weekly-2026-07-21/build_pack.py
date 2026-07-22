from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GREEN = RGBColor(0x00, 0x7B, 0x5F)
BLACK = RGBColor(0x11, 0x11, 0x11)
GREY  = RGBColor(0x55, 0x55, 0x55)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(11); st.font.color.rgb = BLACK

def h1(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=GREEN
    p.space_after=Pt(4); return p
def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=GREEN
    return p
def sub(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic=True; r.font.color.rgb=GREY; r.font.size=Pt(10); return p
def body(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=bold; return p
def field(label, val):
    p = doc.add_paragraph(); r = p.add_run(label+"  "); r.bold=True; r.font.color.rgb=BLACK
    r2 = p.add_run(val); return p

h1("AXRIK — Weekly Outreach Pack")
sub("Prepared 21 July 2026 · Liverpool City Region food & drink · Prep only — nothing has been sent or posted.")

h2("Do this first")
for t in [
    "Give each business's Instagram/Facebook a 30-second look before contacting — confirm they still take orders by hand/DM.",
    "Send the 5 messages below yourself (Facebook / Instagram / website as noted).",
    "Post the 2 social captions with their matching images (post-1.png, post-2.png in this folder).",
    "Mark anyone you contact as 'contacted' in outreach-tracker.csv.",
]:
    doc.add_paragraph(t, style='List Bullet')

h2("This week's 5 prospects + outreach")

prospects = [
 ("1. Bexleys Craft Butchers",
  "Award-winning independent craft butcher (est. 2005) with shops in Old Swan, Tuebrook, Huyton and Carr Farm, Meols.",
  "Old Swan / Huyton / Meols — Liverpool, Knowsley & Wirral",
  "Website bexleys.co.uk · Facebook & Instagram @bexleysbutchers",
  "Multi-shop independent that already takes seasonal/festive orders over the counter and Facebook — the classic JG Foods pattern at slightly bigger scale.",
  "Hi — I'm Phil from AXRIK, a local company that builds simple apps for food businesses. I recently worked with JG Foods, a mobile butcher who took every order by hand through social media. Now his customers order through an app on their phone and his weekly round builds itself. For a few shops taking festive and counter orders over Facebook, an app could pull all of that into one tidy list. Happy to show you a 2-minute look, no pressure — his story's here: axrik.com/jg-foods. Cheers, Phil"),

 ("2. Presly Butchers",
  "Independent local butcher with an active Facebook page for orders and specials.",
  "Knowsley area, Merseyside",
  "Facebook @presly.butchers",
  "Small independent butcher running orders and specials through Facebook — exactly the manual/DM setup an app tidies up.",
  "Hi — I'm Phil from AXRIK, a local firm building simple apps for food businesses. I recently helped JG Foods, a mobile butcher who took every order by hand on social media. Now people order through an app on their phone and his weekly round builds itself. For a butcher taking orders and specials over Facebook, the same idea could save real time each week and stop things slipping through the messages. I'd be glad to show you a quick 2-minute look, no hard sell: axrik.com/jg-foods. Best, Phil"),

 ("3. West Coast Coffee",
  "Independent, family-run small-batch coffee roaster with bag sales, delivery and subscriptions.",
  "Liverpool",
  "Website westcoastcoffee.shop · Instagram",
  "Family-run roaster juggling one-off bag orders and subscriptions — a simple app could handle repeat orders and 'same again' without the DMs.",
  "Hi — I'm Phil from AXRIK. We build straightforward apps for independent food and drink businesses. I recently worked with JG Foods, a mobile butcher who ran every order by hand through social media. Now his customers order through an app and his weekly round sorts itself out. For a roaster juggling bag orders and subscriptions, an app could make repeat orders and 'same again' effortless for your regulars. I'd love to show you a quick 2-minute look — no hard sell: axrik.com/jg-foods. Best, Phil"),

 ("4. Bakery Barn Crosby",
  "Micro bakery in Little Crosby — everything freshly made on site, orders and collections via Instagram.",
  "Little Crosby, Sefton",
  "Instagram @bakerybarncrosby",
  "Small-batch bakery taking orders and collections through Instagram DMs — a neat fit for a tidy order-and-collection app.",
  "Hi — I'm Phil from AXRIK, a local company building simple apps for food businesses. I recently helped JG Foods, a mobile butcher who took every order by hand through social media DMs. Now his customers order through an app and his weekly round builds itself. For a busy little bakery fielding orders and collections over Instagram, the same idea could take a lot off your plate on baking days. Happy to show you a 2-minute look whenever suits, no pressure: axrik.com/jg-foods. Cheers, Phil"),

 ("5. The Farmers Box",
  "Organic fruit & veg (and meat) box supplier delivering direct to homes and businesses in West Lancashire.",
  "Upholland, West Lancashire",
  "Website thefarmersbox.co.uk · Facebook",
  "Box-delivery business built on recurring weekly orders and rounds — the exact 'the round builds itself' story from JG Foods.",
  "Hi — I'm Phil from AXRIK. We build easy little apps for independent food businesses. I recently worked with JG Foods, a mobile butcher who took every order by hand through social media. Now his customers order through an app on their phone and his weekly delivery round builds itself. For a veg-box business run on weekly orders and rounds, that's exactly the bit an app can quietly take care of. I'd be glad to show you a 2-minute look, no pressure: axrik.com/jg-foods. Best, Phil"),
]

for name, what, area, find, why, msg in prospects:
    p = doc.add_paragraph(); r=p.add_run(name); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLACK
    field("What they do:", what)
    field("Area:", area)
    field("Find them:", find)
    field("Why a fit:", why)
    field("Draft message:", msg)
    doc.add_paragraph()

h2("Follow-ups due")
body("None due this week. The 10 prospects already in the tracker were added on 16 July (5 days ago), so they're not yet due a nudge. First follow-ups fall due from the week of 23–30 July — I'll flag them in that run. Tip: aim to send the 16 July batch this week so the nudge timing stays on track.")

h2("Social posts (with images)")

p = doc.add_paragraph(); r=p.add_run("Post 1 — No app-store / cost angle"); r.bold=True; r.font.color.rgb=BLACK
field("Image file:", "post-1.png")
field("Caption:", "A proper app used to mean a proper bill — the kind of quote that ran from £30,000 to £150,000 with an agency, plus the app-store fees to sit on top. That maths never worked for a small food business. AXRIK builds you a real, installable app that lives on your customers' phones — bespoke to how you actually take orders — without the six-figure invoice or the app-store gatekeepers. Same result your customers see, a fraction of the cost, and it's yours. Curious what that looks like for your shop? Take a look: axrik.com/jg-foods")
doc.add_paragraph()

p = doc.add_paragraph(); r=p.add_run("Post 2 — JG Foods case-study spotlight"); r.bold=True; r.font.color.rgb=BLACK
field("Image file:", "post-2.png")
field("Caption:", "Not long ago, JG Foods took every single order by hand — messages pinging in across social media, all copied out one by one. It worked, but it ate the evenings. So we built them a simple app. Now customers order straight from their phones, and the weekly delivery round more or less builds itself from what's come in. Same friendly butcher, none of the admin scramble. That's the whole idea behind AXRIK: quiet, useful tech that gives a small food business its time back. See how it worked: axrik.com/jg-foods")
doc.add_paragraph()

sub("Prep only — nothing has been sent or posted. Review, then post and send yourself.")

doc.save("AXRIK-Outreach-Pack-2026-07-21.docx")
print("saved docx")
